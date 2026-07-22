import json,sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
p=json.load(open(sys.argv[1],encoding='utf-8'))
text=p['plainText']; docx_path=p['docxPath']; pdf_path=p['pdfPath']
lines=text.splitlines(); name=lines[0] if lines else ''; contact=lines[1] if len(lines)>1 else ''; title=lines[2] if len(lines)>2 else ''
heads={'PROFESSIONAL SUMMARY','CORE SKILLS','PROFESSIONAL EXPERIENCE'}; sections=[]; cur=None
for line in lines[3:]:
 t=line.strip()
 if t in heads: cur=[t,[]]; sections.append(cur)
 elif t and cur: cur[1].append(t)
d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(.5)
for txt,size,bold in [(name,17,True),(contact,9.5,False),(title,11,True)]:
 q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run(txt); r.bold=bold; r.font.size=Pt(size)
for h,items in sections:
 q=d.add_paragraph(); r=q.add_run(h); r.bold=True; r.font.size=Pt(11)
 for item in items:
  bullet=item.startswith('• '); q=d.add_paragraph(style='List Bullet' if bullet else None); q.add_run(item[2:] if bullet else item).font.size=Pt(10)
d.core_properties.title=title; d.core_properties.author='Life Vault HQ Resume Product'; d.save(docx_path)
styles=getSampleStyleSheet(); styles.add(ParagraphStyle(name='CenterSmall',parent=styles['Normal'],alignment=TA_CENTER,fontSize=9.5,leading=11)); styles.add(ParagraphStyle(name='CenterTitle',parent=styles['Heading1'],alignment=TA_CENTER,fontSize=17,leading=20)); styles.add(ParagraphStyle(name='Section',parent=styles['Heading2'],fontSize=11,leading=13,spaceBefore=8,spaceAfter=4)); styles['Normal'].fontSize=9.5; styles['Normal'].leading=12
story=[Paragraph(name,styles['CenterTitle']),Paragraph(contact,styles['CenterSmall']),Paragraph(title,styles['CenterSmall']),Spacer(1,8)]
for h,items in sections:
 story.append(Paragraph(h,styles['Section']))
 for item in items:
  if item.startswith('• '): story.append(ListFlowable([ListItem(Paragraph(item[2:],styles['Normal']))],bulletType='bullet',leftIndent=14))
  else: story.append(Paragraph(item,styles['Normal']))
SimpleDocTemplate(pdf_path,pagesize=letter,rightMargin=36,leftMargin=36,topMargin=30,bottomMargin=30,title=title,author='Life Vault HQ Resume Product').build(story)
